import os
import docx
from typing import List
from docx.shared import Pt, Inches, Cm
from typing import List
import pandas as pd


def change_orientation(document, orientation: str):
    """
    Change the orientation of the document.

    :param document: The document object.
    :type document: docx.Document
    :param orientation: The desired orientation of the document. Options are 'landscape' or 'portrait'.
    :type orientation: str

    :raises ValueError: If the provided `orientation` is not one of 'landscape' or 'portrait'.
    """
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    if orientation.lower() == 'landscape':
        current_section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    elif orientation.lower() == 'portrait':
        current_section.orientation = docx.enum.section.WD_ORIENT.PORTRAIT
    else:
        raise ValueError('Specify "landscape" or "portrait" for the orientation.')
    current_section.page_width = new_width


def format_cell(cell, **kwargs):
    """
    Format a cell in a table.

    :param cell: The cell to be formatted.
    :type cell: Cell
    :param kwargs: Additional keyword arguments.
    :type kwargs: dict
    :keyword font_size: The font size for the cell. Default is None.
    :type font_size: int
    :keyword bold: Whether the cell text should be bold. Default is None.
    :type bold: bool
    :keyword paragraph_font_size: The font size for the paragraph in the cell. Default is 11.
    :type paragraph_font_size: int
    :keyword table_font_size: The font size for the table. Default is 10.
    :type table_font_size: int
    :keyword cell_alignment: The alignment of the cell text. Options are 'left', 'center', or 'right'.
                            Default is 'left'.
    :type cell_alignment: str

    :raises ValueError: If the provided `cell_alignment` is not one of 'left', 'center', or 'right'.
    """

    # Parse keyword arguments
    bold = kwargs.get('bold', None)
    paragraph_font_size = kwargs.get('paragraph_font_size', 11)
    table_font_size = kwargs.get('table_font_size', 10)
    alignment = kwargs.get('alignment', 'left')

    # Get pointer
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0]

    # Format the cell
    run.font.size = Pt(table_font_size)
    if bold:
        run.bold = True

    p = cell.paragraphs[0]
    if alignment.lower() == 'left':
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    elif alignment.lower() == 'center':
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    elif alignment.lower() == 'right':
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    else:
        raise ValueError('Alignment must be left, center, or right.')


def set_col_widths(table, widths):
    """
    Set the widths of columns in a table.

    :param table: The table object.
    :type table: Table
    :param widths: A list of column widths.
    :type widths: list
    """

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def generate_report(data_frames: List[pd.DataFrame], **kwargs) -> docx.Document:
    """
    Generate a water quality model report with summary statistics using python-docx.

    :param data_frames: List of Pandas DataFrames containing water quality data.
    :type data_frames: List[pd.DataFrame]
    :param **kwargs: Additional keyword arguments for model_name, location, and time_period.
    :type **kwargs: Any
    :return: The generated report document.
    :rtype: docx.Document
    """

    # Assign keyword arguments to variables
    model_name = kwargs.get('model_name', None)
    location = kwargs.get('location', None)
    time_period = kwargs.get('time_period', None)
    table_font_size = kwargs.get('table_font_size', 10)

    # Create a new document
    doc = docx.Document()

    # Change to landscape orientation
    change_orientation(doc, "landscape")

    # Add the report title
    doc.add_heading("Water Quality Model Report", level=0)

    # Add model information
    doc.add_heading("Model Information", level=1)
    if model_name:
        doc.add_paragraph(f"Model Name: {model_name}")
    if location:
        doc.add_paragraph(f"Location: {location}")
    if time_period:
        doc.add_paragraph(f"Time Period: {time_period}")

    # Add data summary
    doc.add_heading("Data Summary", level=1)
    doc.add_paragraph("Summary Statistics:")

    for df in data_frames:
        # Compute summary statistics for the current dataframe
        summary_statistics = df.describe().to_dict()
        col1 = df.columns[0]
        dict1 = summary_statistics[col1]
        keys = dict1.keys()

        # Add sub-heading for data filename
        file_path = df.attrs['Filename']
        directory, filename = os.path.split(file_path)
        doc.add_heading(filename, level=2)

        # Create a table for summary statistics
        table = doc.add_table(rows=1, cols=len(df.columns) + 1)
        # table.autofit = True
        table.style = "Table Grid"

        # Add table headers
        table_header_cells = table.rows[0].cells
        cell = table_header_cells[0]
        cell.text = "Statistic"
        format_cell(cell, table_font_size=table_font_size, bold=True, alignment='left')
        for j, col in enumerate(df.columns):
            cell = table_header_cells[j + 1]
            cell.text = col
            format_cell(cell, table_font_size=table_font_size, bold=True, alignment='center')

        # Populate tabble
        for i, key in enumerate(keys):
            row = table.add_row().cells
            cell = row[0]
            cell.text = key
            format_cell(cell, table_font_size=table_font_size, bold=False, alignment='left')
            for j, col in enumerate(df.columns):
                d = summary_statistics[col][key]
                cell = row[j + 1]
                cell.text = f'{d:.1f}'
                format_cell(cell, table_font_size=table_font_size, bold=False, alignment='right')

        widths = [Inches(0.1)]
        for col in df.columns:
            widths.append(Inches(0.1))
        set_col_widths(table, widths)

    return doc